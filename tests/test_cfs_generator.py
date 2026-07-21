from io import BytesIO
from zipfile import ZipFile

from docx import Document

from Contracts.generators.cfs_generator import (
    BLANK_CFS_CONTEXT,
    CFS_ANNEX_SECTION_HEADINGS,
    CFS_MAIN_SECTION_HEADINGS,
    generate_blank_cfs_docx,
)


def test_blank_cfs_form_has_writing_lines_and_no_template_markers() -> None:
    output = generate_blank_cfs_docx().getvalue()

    with ZipFile(BytesIO(output)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")

    assert "{{" not in document_xml
    assert "{%" not in document_xml

    document = Document(BytesIO(output))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    for writing_line in BLANK_CFS_CONTEXT.values():
        assert writing_line in text

    assert "Name: ________________________________________" in text
    assert "Date: ____________________" in text


def test_cfs_sections_have_stable_pagination_controls() -> None:
    document = Document(BytesIO(generate_blank_cfs_docx().getvalue()))
    paragraphs = document.paragraphs
    paragraph_text = [" ".join(paragraph.text.split()) for paragraph in paragraphs]

    signature_index = paragraph_text.index("SIGNATURES")
    annex_index = paragraph_text.index("Annex A – Scope of Services")
    main_indices = [paragraph_text.index(heading) for heading in CFS_MAIN_SECTION_HEADINGS]
    annex_indices = [paragraph_text.index(heading) for heading in CFS_ANNEX_SECTION_HEADINGS]

    for start_index, end_index in zip(
        main_indices,
        main_indices[1:] + [signature_index],
    ):
        meaningful = [
            index for index in range(start_index, end_index) if paragraph_text[index]
        ]
        assert meaningful
        assert all(
            paragraphs[index].paragraph_format.keep_with_next is True
            for index in range(start_index, meaningful[-1])
        )

    for start_index, end_index in zip(
        annex_indices,
        annex_indices[1:] + [len(paragraphs)],
    ):
        meaningful = [
            index for index in range(start_index, end_index) if paragraph_text[index]
        ]
        assert meaningful
        assert all(
            paragraphs[index].paragraph_format.keep_with_next is True
            for index in range(start_index, meaningful[-1])
        )

    assert all(
        paragraph.paragraph_format.keep_together is True
        for paragraph in paragraphs
    )
    assert all(
        paragraph.paragraph_format.widow_control is True
        for paragraph in paragraphs
    )
    assert paragraphs[annex_index].paragraph_format.page_break_before is True
    assert not paragraphs[annex_index]._p.xpath(".//w:br[@w:type='page']")


def test_entire_agreement_heading_stays_with_its_body() -> None:
    document = Document(BytesIO(generate_blank_cfs_docx().getvalue()))
    paragraph_text = [" ".join(paragraph.text.split()) for paragraph in document.paragraphs]
    heading_index = paragraph_text.index("Entire Agreement")

    assert document.paragraphs[heading_index].paragraph_format.keep_with_next is True
    assert document.paragraphs[heading_index + 1].paragraph_format.keep_together is True
